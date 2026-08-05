#!/usr/bin/env python3
"""Render constrained resume or interview-prep Markdown to DOCX and PDF.

The Markdown file remains the reader-facing source. This script applies a fixed
A4 single-column style, converts DOCX to PDF with LibreOffice, renders PDF QA
images with Poppler, and reports estimated/actual pages plus text parity.
"""

from __future__ import annotations

import argparse
import difflib
import html
import json
import math
import os
from pathlib import Path
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
from dataclasses import dataclass
from typing import Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, RGBColor
from docx.text.run import Run
from pypdf import PdfReader


WESTERN_FONT = os.environ.get("TAILOR_RESUME_WESTERN_FONT", "Arial")


def default_cjk_font() -> str:
    if sys.platform == "darwin":
        return "Hiragino Sans GB"
    if sys.platform.startswith("win"):
        return "Microsoft YaHei"
    return "Noto Sans CJK SC"


CJK_FONT = os.environ.get("TAILOR_RESUME_CJK_FONT", default_cjk_font())
BODY_SIZE_PT = 10.5
BODY_LINE_SPACING = 1.08
PAGE_MARGIN_MM = 18
PARITY_THRESHOLD = 0.97

HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^\s*[-+*]\s+(.+?)\s*$")
ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+?)\s*$")
TABLE_DIVIDER_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$",
    re.MULTILINE,
)
INLINE_RE = re.compile(
    r"(\[[^\]]+\]\([^)]+\)|\*\*.+?\*\*|__.+?__|`[^`]+`|"
    r"(?<!\*)\*[^*\n]+\*(?!\*)|(?<!_)_[^_\n]+_(?!_))"
)
LINK_RE = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")
NON_ASCII_DASHES = str.maketrans(
    {
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2212": "-",
    }
)


@dataclass(frozen=True)
class Block:
    kind: str
    text: str = ""
    level: int = 0


def visible_link(label: str, url: str) -> str:
    label = label.strip()
    url = url.strip()
    return label or url.removeprefix("mailto:")


def strip_inline(text: str) -> str:
    def replace(match: re.Match[str]) -> str:
        token = match.group(0)
        link = LINK_RE.match(token)
        if link:
            return visible_link(link.group(1), link.group(2))
        if token.startswith(("**", "__")):
            return token[2:-2]
        if token.startswith(("`", "*", "_")):
            return token[1:-1]
        return token

    return html.unescape(INLINE_RE.sub(replace, text)).translate(NON_ASCII_DASHES)


def parse_markdown(markdown: str) -> list[Block]:
    if TABLE_DIVIDER_RE.search(markdown):
        raise ValueError("Markdown tables are not supported in ATS resume content.")

    blocks: list[Block] = []
    paragraph_lines: list[str] = []
    in_comment = False

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append(Block("paragraph", " ".join(paragraph_lines).strip()))
            paragraph_lines.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if in_comment:
            if "-->" in stripped:
                in_comment = False
            continue
        if stripped.startswith("<!--"):
            if "-->" not in stripped:
                in_comment = True
            continue

        if not stripped:
            flush_paragraph()
            continue

        if stripped in {"---", "***", "___"}:
            flush_paragraph()
            blocks.append(Block("rule"))
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            flush_paragraph()
            blocks.append(Block("heading", heading.group(2), len(heading.group(1))))
            continue

        bullet = BULLET_RE.match(line)
        if bullet:
            flush_paragraph()
            blocks.append(Block("bullet", bullet.group(1)))
            continue

        ordered = ORDERED_RE.match(line)
        if ordered:
            flush_paragraph()
            blocks.append(Block("ordered", ordered.group(1)))
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            blocks.append(Block("quote", stripped[2:].strip()))
            continue

        paragraph_lines.append(stripped)

    flush_paragraph()

    h1_count = sum(block.kind == "heading" and block.level == 1 for block in blocks)
    h2_count = sum(block.kind == "heading" and block.level == 2 for block in blocks)
    if h1_count != 1:
        raise ValueError(
            f"Document Markdown must contain exactly one H1 title; found {h1_count}."
        )
    if h2_count == 0:
        raise ValueError("Document Markdown must contain at least one H2 section.")
    return blocks


def contains_cjk(text: str) -> bool:
    return any(
        "\u3400" <= char <= "\u9fff"
        or "\uf900" <= char <= "\ufaff"
        or "\u3040" <= char <= "\u30ff"
        or "\uac00" <= char <= "\ud7af"
        for char in text
    )


def set_run_font(run, size: float | None = None) -> None:
    primary_font = CJK_FONT if contains_cjk(run.text) else WESTERN_FONT
    run.font.name = primary_font
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), primary_font)
    rfonts.set(qn("w:hAnsi"), primary_font)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rfonts.set(qn("w:cs"), primary_font)
    if contains_cjk(run.text):
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:eastAsia"), "zh-CN")


def set_style_font(style, size: float, bold: bool = False, color: str = "111111") -> None:
    style.font.name = WESTERN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), WESTERN_FONT)
    rfonts.set(qn("w:hAnsi"), WESTERN_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rfonts.set(qn("w:cs"), WESTERN_FONT)


def ensure_paragraph_style(document: Document, name: str):
    styles = document.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    set_style_font(normal, BODY_SIZE_PT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING
    normal.paragraph_format.widow_control = True

    resume_name = ensure_paragraph_style(document, "Resume Name")
    set_style_font(resume_name, 18, bold=True, color="000000")
    resume_name.paragraph_format.space_before = Pt(0)
    resume_name.paragraph_format.space_after = Pt(2)
    resume_name.paragraph_format.line_spacing = 1.0
    resume_name.paragraph_format.keep_with_next = True

    contact = ensure_paragraph_style(document, "Resume Contact")
    set_style_font(contact, 9.5, color="333333")
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(4)
    contact.paragraph_format.line_spacing = 1.0
    contact.paragraph_format.keep_with_next = True

    heading1 = document.styles["Heading 1"]
    set_style_font(heading1, 12.5, bold=True, color="000000")
    heading1.paragraph_format.space_before = Pt(6)
    heading1.paragraph_format.space_after = Pt(2)
    heading1.paragraph_format.line_spacing = 1.0
    heading1.paragraph_format.keep_with_next = True

    heading2 = document.styles["Heading 2"]
    set_style_font(heading2, 11, bold=True, color="111111")
    heading2.paragraph_format.space_before = Pt(4)
    heading2.paragraph_format.space_after = Pt(0)
    heading2.paragraph_format.line_spacing = 1.0
    heading2.paragraph_format.keep_with_next = True

    heading3 = document.styles["Heading 3"]
    set_style_font(heading3, 10.5, bold=True, color="111111")
    heading3.paragraph_format.space_before = Pt(3)
    heading3.paragraph_format.space_after = Pt(0)
    heading3.paragraph_format.line_spacing = 1.0
    heading3.paragraph_format.keep_with_next = True

    metadata = ensure_paragraph_style(document, "Resume Metadata")
    set_style_font(metadata, 9.5, color="444444")
    metadata.font.italic = True
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(1)
    metadata.paragraph_format.line_spacing = 1.0
    metadata.paragraph_format.keep_with_next = True

    quote = ensure_paragraph_style(document, "Resume Note")
    set_style_font(quote, 9.5, color="444444")
    quote.font.italic = True
    quote.paragraph_format.left_indent = Mm(4)
    quote.paragraph_format.space_after = Pt(2)

    for list_name in ("List Bullet", "List Number"):
        list_style = document.styles[list_name]
        set_style_font(list_style, BODY_SIZE_PT)
        list_style.paragraph_format.left_indent = Mm(5)
        list_style.paragraph_format.first_line_indent = Mm(-3)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(1)
        list_style.paragraph_format.line_spacing = BODY_LINE_SPACING
        list_style.paragraph_format.widow_control = True


def add_inline_runs(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(
                html.unescape(text[cursor : match.start()]).translate(NON_ASCII_DASHES)
            )
            set_run_font(run)

        token = match.group(0)
        link = LINK_RE.match(token)
        if link:
            add_hyperlink(paragraph, link.group(1), link.group(2))
            cursor = match.end()
            continue
        elif token.startswith(("**", "__")):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        set_run_font(run)
        cursor = match.end()

    if cursor < len(text):
        run = paragraph.add_run(html.unescape(text[cursor:]).translate(NON_ASCII_DASHES))
        set_run_font(run)


def add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run_element = OxmlElement("w:r")
    hyperlink.append(run_element)
    run = Run(run_element, paragraph)
    run.text = visible_link(label, url)
    set_run_font(run)
    run.font.underline = True
    paragraph._p.append(hyperlink)


def add_rule(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(2)
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7B7B7")
    borders.append(bottom)
    ppr.append(borders)


def build_docx(
    blocks: Iterable[Block], output_path: Path, document_kind: str = "resume"
) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(PAGE_MARGIN_MM)
    section.bottom_margin = Mm(PAGE_MARGIN_MM)
    section.left_margin = Mm(PAGE_MARGIN_MM)
    section.right_margin = Mm(PAGE_MARGIN_MM)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    configure_styles(document)
    if document_kind == "interview":
        document.core_properties.title = "Interview Preparation"
        document.core_properties.subject = "Job-specific interview preparation"
    else:
        document.core_properties.title = "Tailored Resume"
        document.core_properties.subject = "Job-specific resume"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = ""

    seen_name = False
    seen_section = False
    contact_written = False

    for block in blocks:
        if block.kind == "rule":
            add_rule(document.add_paragraph())
            continue

        if block.kind == "heading":
            if block.level == 1:
                paragraph = document.add_paragraph(style="Resume Name")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                seen_name = True
            elif block.level == 2:
                paragraph = document.add_paragraph(style="Heading 1")
                seen_section = True
            elif block.level == 3:
                paragraph = document.add_paragraph(style="Heading 2")
            else:
                paragraph = document.add_paragraph(style="Heading 3")
            add_inline_runs(paragraph, block.text)
            continue

        if block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.keep_together = True
            add_inline_runs(paragraph, block.text)
            continue

        if block.kind == "ordered":
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.keep_together = True
            add_inline_runs(paragraph, block.text)
            continue

        if block.kind == "quote":
            paragraph = document.add_paragraph(style="Resume Note")
            add_inline_runs(paragraph, block.text)
            continue

        text = block.text
        is_metadata = (
            len(text) >= 2
            and ((text.startswith("*") and text.endswith("*")) or
                 (text.startswith("_") and text.endswith("_")))
            and not text.startswith(("**", "__"))
        )
        if is_metadata:
            paragraph = document.add_paragraph(style="Resume Metadata")
            add_inline_runs(paragraph, text[1:-1])
        elif seen_name and not seen_section and not contact_written:
            paragraph = document.add_paragraph(style="Resume Contact")
            add_inline_runs(paragraph, text)
            contact_written = True
        else:
            paragraph = document.add_paragraph(style="Normal")
            add_inline_runs(paragraph, text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def display_width(text: str) -> int:
    width = 0
    for char in strip_inline(text):
        if char == "\t":
            width += 4
        elif unicodedata.east_asian_width(char) in {"W", "F"}:
            width += 2
        else:
            width += 1
    return width


def estimate_pages(blocks: Iterable[Block]) -> tuple[float, int]:
    equivalent_lines = 0.0
    seen_name = False
    seen_section = False
    contact_counted = False
    for block in blocks:
        width = max(1, display_width(block.text))
        if block.kind == "heading":
            if block.level == 1:
                equivalent_lines += math.ceil(width / 62) * 1.8 + 0.5
                seen_name = True
            elif block.level == 2:
                equivalent_lines += math.ceil(width / 78) * 1.25 + 0.7
                seen_section = True
            else:
                equivalent_lines += math.ceil(width / 84) * 1.1 + 0.25
        elif block.kind in {"bullet", "ordered"}:
            equivalent_lines += math.ceil(width / 88) + 0.08
        elif block.kind == "rule":
            equivalent_lines += 0.5
        elif block.kind == "quote":
            equivalent_lines += math.ceil(width / 88) + 0.2
        elif seen_name and not seen_section and not contact_counted:
            equivalent_lines += math.ceil(width / 100) * 0.9 + 0.5
            contact_counted = True
        else:
            equivalent_lines += math.ceil(width / 94) + 0.15

    pages = max(1, math.ceil(equivalent_lines / 58.0))
    return round(equivalent_lines, 2), pages


def runtime_dependency_root() -> Path | None:
    executable = Path(sys.executable).resolve()
    for parent in executable.parents:
        if parent.name == "dependencies":
            return parent
    return None


def find_executable(name: str) -> str | None:
    found = shutil.which(name)
    if found:
        return found

    root = runtime_dependency_root()
    if root:
        for directory in (root / "bin" / "override", root / "bin" / "fallback"):
            candidate = directory / name
            if candidate.exists():
                return str(candidate)

    if name == "soffice":
        for candidate in (
            Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
            Path("/usr/local/bin/soffice"),
            Path("/opt/homebrew/bin/soffice"),
        ):
            if candidate.exists():
                return str(candidate)
    return None


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    soffice = find_executable("soffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice/soffice was not found; cannot create the required PDF."
        )

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="tailor-resume-lo-") as temp_name:
        temp_dir = Path(temp_name)
        profile_dir = temp_dir / "profile"
        profile_dir.mkdir()
        converted_dir = temp_dir / "converted"
        converted_dir.mkdir()

        env = os.environ.copy()
        env["HOME"] = str(temp_dir)
        env["TMPDIR"] = "/private/tmp" if Path("/private/tmp").exists() else str(temp_dir)
        user_home = Path.home()
        windows_font_directory = (
            Path(os.environ["WINDIR"]) / "Fonts" if os.environ.get("WINDIR") else None
        )
        font_directories = [
            path
            for path in (
                Path("/System/Library/Fonts"),
                Path("/System/Library/Fonts/Supplemental"),
                Path("/Library/Fonts"),
                user_home / "Library" / "Fonts",
                user_home / ".fonts",
                user_home / ".local" / "share" / "fonts",
                Path("/usr/share/fonts"),
                Path("/usr/local/share/fonts"),
                windows_font_directory,
            )
            if path is not None and path.exists()
        ]
        if font_directories:
            env["SAL_FONTPATH"] = os.pathsep.join(str(path) for path in font_directories)
            font_cache = temp_dir / "font-cache"
            font_cache.mkdir()
            fontconfig_file = temp_dir / "fonts.conf"
            fontconfig_file.write_text(
                "<?xml version=\"1.0\"?>\n"
                "<!DOCTYPE fontconfig SYSTEM \"fonts.dtd\">\n"
                "<fontconfig>\n"
                + "".join(
                    f"  <dir>{html.escape(str(path))}</dir>\n"
                    for path in font_directories
                )
                + f"  <cachedir>{html.escape(str(font_cache))}</cachedir>\n"
                "  <config><rescan><int>0</int></rescan></config>\n"
                "</fontconfig>\n",
                encoding="utf-8",
            )
            env["FONTCONFIG_FILE"] = str(fontconfig_file)
            env["FONTCONFIG_PATH"] = str(temp_dir)
        command = [
            soffice,
            "--headless",
            f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(converted_dir),
            str(docx_path.resolve()),
        ]
        result = subprocess.run(
            command,
            env=env,
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        generated = converted_dir / f"{docx_path.stem}.pdf"
        if result.returncode != 0 or not generated.exists() or generated.stat().st_size == 0:
            details = (result.stderr or result.stdout or "unknown conversion error").strip()
            raise RuntimeError(f"DOCX to PDF conversion failed: {details}")

        if pdf_path.exists():
            pdf_path.unlink()
        shutil.move(str(generated), str(pdf_path))


def render_pdf_pages(pdf_path: Path, qa_dir: Path) -> list[Path]:
    pdftoppm = find_executable("pdftoppm")
    if not pdftoppm:
        raise RuntimeError("pdftoppm was not found; cannot create visual-QA images.")

    qa_dir.mkdir(parents=True, exist_ok=True)
    for existing in qa_dir.glob("page-*.png"):
        existing.unlink()

    prefix = qa_dir / "page"
    result = subprocess.run(
        [pdftoppm, "-png", "-r", "144", str(pdf_path), str(prefix)],
        capture_output=True,
        text=True,
        timeout=120,
        check=False,
    )
    if result.returncode != 0:
        details = (result.stderr or result.stdout or "unknown rendering error").strip()
        raise RuntimeError(f"PDF rendering failed: {details}")

    def page_number(path: Path) -> int:
        match = re.search(r"-(\d+)\.png$", path.name)
        return int(match.group(1)) if match else 0

    images = sorted(qa_dir.glob("page-*.png"), key=page_number)
    if not images:
        raise RuntimeError("PDF rendering completed without producing page images.")
    return images


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def parity_key(text: str) -> str:
    text = text.translate(NON_ASCII_DASHES).lower()
    return "".join(char for char in text if char.isalnum())


def text_parity(left: str, right: str) -> float:
    left_key = parity_key(left)
    right_key = parity_key(right)
    if not left_key or not right_key:
        return 0.0
    return difflib.SequenceMatcher(None, left_key, right_key).ratio()


def render(args: argparse.Namespace) -> dict[str, object]:
    input_path = args.input.resolve()
    docx_path = args.docx.resolve()
    pdf_path = args.pdf.resolve()

    markdown = input_path.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown)
    equivalent_lines, estimated_pages = estimate_pages(blocks)

    build_docx(blocks, docx_path, args.document_kind)
    convert_docx_to_pdf(docx_path, pdf_path)

    reader = PdfReader(str(pdf_path))
    actual_pages = len(reader.pages)
    qa_images: list[Path] = []
    if args.qa_dir:
        qa_images = render_pdf_pages(pdf_path, args.qa_dir.resolve())

    docx_text = extract_docx_text(docx_path)
    pdf_text = extract_pdf_text(pdf_path)
    parity = text_parity(docx_text, pdf_text)

    report: dict[str, object] = {
        "input_markdown": str(input_path),
        "document_kind": args.document_kind,
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "estimated_equivalent_lines": equivalent_lines,
        "estimated_pages": estimated_pages,
        "actual_pdf_pages": actual_pages,
        "target_pages": args.target_pages,
        "page_cap_exceeded": (
            args.target_pages is not None and actual_pages > args.target_pages
        ),
        "docx_pdf_text_parity": round(parity, 4),
        "minimum_text_parity": args.min_parity,
        "text_parity_passed": parity >= args.min_parity,
        "qa_images": [str(path.resolve()) for path in qa_images],
        "visual_qa_required": bool(args.qa_dir),
        "visual_qa_completed": False,
        "layout": {
            "page": "A4",
            "margins_mm": PAGE_MARGIN_MM,
            "body_font_western": WESTERN_FONT,
            "body_font_cjk": CJK_FONT,
            "body_size_pt": BODY_SIZE_PT,
            "line_spacing": BODY_LINE_SPACING,
            "columns": 1,
        },
    }

    if args.report:
        args.report.resolve().parent.mkdir(parents=True, exist_ok=True)
        args.report.resolve().write_text(
            json.dumps(report, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
    return report


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Render resume or interview-prep Markdown to matching DOCX/PDF files."
    )
    parser.add_argument("input", type=Path, help="Reader-facing Markdown source file")
    parser.add_argument(
        "--document-kind",
        choices=["resume", "interview"],
        default="resume",
        help="Document metadata profile (default: resume)",
    )
    parser.add_argument("--docx", type=Path, help="Output DOCX path")
    parser.add_argument("--pdf", type=Path, help="Output PDF path")
    parser.add_argument(
        "--qa-dir",
        type=Path,
        help="Temporary directory for page-*.png visual-QA images",
    )
    parser.add_argument(
        "--target-pages",
        type=int,
        choices=range(1, 5),
        metavar="{1,2,3,4}",
        help="Maximum allowed final PDF pages",
    )
    parser.add_argument(
        "--min-parity",
        type=float,
        default=PARITY_THRESHOLD,
        help=f"Minimum DOCX/PDF normalized text parity (default: {PARITY_THRESHOLD})",
    )
    parser.add_argument("--report", type=Path, help="Optional JSON render report path")
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    if not args.input.exists():
        parser.error(f"input does not exist: {args.input}")
    if args.docx is None:
        args.docx = args.input.with_suffix(".docx")
    if args.pdf is None:
        args.pdf = args.input.with_suffix(".pdf")
    if args.min_parity < 0 or args.min_parity > 1:
        parser.error("--min-parity must be between 0 and 1")

    try:
        report = render(args)
    except Exception as error:
        print(json.dumps({"error": str(error)}, ensure_ascii=False, indent=2))
        return 1

    print(json.dumps(report, ensure_ascii=False, indent=2))
    if report["page_cap_exceeded"]:
        return 2
    if not report["text_parity_passed"]:
        return 3
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
